#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
import_docx.py – Trích xuất 379 câu hỏi HVAC_ASHRAE_VSCD_2026 từ 2 file Word
                  → data/questions.json + data/import-report.json

Cách dùng:
    cd <project_root>
    pip install -r scripts/requirements.txt
    python scripts/import_docx.py
"""

import sys
import io
import os
import re
import json
import hashlib
import datetime
import traceback
from pathlib import Path

# Force UTF-8 stdout for Windows
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: Thư viện python-docx chưa được cài. Chạy:")
    print("  pip install -r scripts/requirements.txt")
    sys.exit(1)

# ─── Paths ────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DOCX_NAME    = "HVAC_ASHRAE_VSCD_2026_Question_Bank_Full_Bilingual_EN_VI.docx"
DOCX_PATH    = PROJECT_ROOT / DOCX_NAME
DOCX2_NAME   = "Bo_75_cau_trac_nghiem_HVAC_song_ngu_Anh_Viet_ASHRAE_62_1_2022.docx"
DOCX2_PATH   = PROJECT_ROOT / DOCX2_NAME
DOCX3_NAME   = "Bo_60_cau_trac_nghiem_ASHRAE_52_2_2017_Bloom_Song_ngu_EN_VI.docx"
DOCX3_PATH   = PROJECT_ROOT / DOCX3_NAME

DATA_DIR       = PROJECT_ROOT / "data"
IMG_DIR        = PROJECT_ROOT / "public" / "assets" / "questions"
QUESTIONS_JSON = DATA_DIR / "questions.json"
REPORT_JSON    = DATA_DIR / "import-report.json"

EXPECTED_COUNT = 439

# ─── Validate DOCX exists ─────────────────────────────────────────────────────
if not DOCX_PATH.exists():
    print(f"ERROR: Không tìm thấy file Word '{DOCX_NAME}' trong {PROJECT_ROOT}")
    sys.exit(1)
if not DOCX2_PATH.exists():
    print(f"ERROR: Không tìm thấy file Word '{DOCX2_NAME}' trong {PROJECT_ROOT}")
    sys.exit(1)
if not DOCX3_PATH.exists():
    print(f"ERROR: Không tìm thấy file Word '{DOCX3_NAME}' trong {PROJECT_ROOT}")
    sys.exit(1)

DATA_DIR.mkdir(parents=True, exist_ok=True)
IMG_DIR.mkdir(parents=True, exist_ok=True)

# ─── SHA-256 ──────────────────────────────────────────────────────────────────
def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

# ─── Regex ────────────────────────────────────────────────────────────────────
QH_RE = re.compile(
    r"^(Q\d{1,4})\s*[·\-]\s*(ASHRAE\s+[\d.]+(?:-\d+)?(?:\s+\w+)*|[\w.\s]+?)\s*[·\-]\s*(.+)$",
    re.IGNORECASE
)

AH_RE = re.compile(
    r"^(Q\d{1,4})\s*[·\-]\s*(?:ĐÁP\s*ÁN|CORRECT\s*ANSWER|ANSWER)\s+([A-D])\s*[·\-]?\s*(.*)$",
    re.IGNORECASE
)

OPT_RE = re.compile(r"^([A-D])[.)]\s+(.+)$", re.DOTALL)
CORRECT_TEXT_RE = re.compile(r"(?:correct\s*answer|đáp\s*án\s*đúng)[:\s]+(.+)", re.IGNORECASE | re.DOTALL)
EXPL_RE = re.compile(r"^(?:giải\s*thích|explanation)[:\s]+(.+)$", re.IGNORECASE | re.DOTALL)
SRC_RE = re.compile(r"^(?:nguồn|source)[^:]*:\s*(.+)$", re.IGNORECASE)

BLOOM_H_RE = re.compile(r"BLOOM\s+(\d)\s*[·\-]\s*(REMEMBER|UNDERSTAND|APPLY|ANALYZE|EVALUATE|CREATE)", re.IGNORECASE)
STD_H_RE   = re.compile(r"(ASHRAE\s+[\d.]+(?:-\d+)?)", re.IGNORECASE)

# ─── Extract images ───────────────────────────────────────────────────────────
def extract_all_images(doc) -> tuple:
    img_paths = {}
    img_blobs = {}
    img_counter = 0
    rels = doc.part.rels
    ext_map = {
        "image/png": "png", "image/jpeg": "jpg", "image/gif": "gif",
        "image/bmp": "bmp", "image/tiff": "tif", "image/wmf": "wmf", "image/emf": "emf",
    }
    for rId, rel in rels.items():
        if "image" in rel.reltype.lower():
            img_counter += 1
            try:
                img_part = rel.target_part
                ext = ext_map.get(img_part.content_type, "png")
                fname = f"img_{img_counter:04d}.{ext}"
                web_path = f"/assets/questions/{fname}"
                img_paths[rId] = web_path
                img_blobs[web_path] = (fname, img_part.blob)
            except Exception as e:
                print(f"  WARNING img rId={rId}: {e}")
    return img_paths, img_blobs

def extract_all_images_ashrae62(doc, start_num=14) -> tuple:
    img_paths = {}
    img_blobs = {}
    img_counter = start_num - 1
    rels = doc.part.rels
    ext_map = {
        "image/png": "png", "image/jpeg": "jpg", "image/gif": "gif",
        "image/bmp": "bmp", "image/tiff": "tif", "image/wmf": "wmf", "image/emf": "emf",
    }
    for rId in sorted(rels.keys(), key=lambda x: int(x[3:]) if x[3:].isdigit() else x):
        rel = rels[rId]
        if "image" in rel.reltype.lower():
            img_counter += 1
            try:
                img_part = rel.target_part
                ext = ext_map.get(img_part.content_type, "png")
                fname = f"img_{img_counter:04d}.{ext}"
                web_path = f"/assets/questions/{fname}"
                img_paths[rId] = web_path
                img_blobs[web_path] = (fname, img_part.blob)
            except Exception as e:
                print(f"  WARNING img2 rId={rId}: {e}")
    return img_paths, img_blobs

def get_para_image_rids(para) -> list:
    rids = []
    for elem in para._p.iter():
        if elem.tag == qn("a:blip"):
            rid = elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if rid: rids.append(rid)
        if "imagedata" in elem.tag.lower():
            rid = elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rid: rids.append(rid)
    return rids

def para_full_text(para) -> str:
    parts = []
    for elem in para._p.iter():
        if elem.tag == qn("w:t"):
            parts.append(elem.text or "")
        elif elem.tag == qn("w:br"):
            parts.append("\n")
        elif elem.tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)

# ─── Option parser ────────────────────────────────────────────────────────────
def parse_option(raw_text: str) -> dict:
    raw_text = raw_text.strip()
    m = OPT_RE.match(raw_text)
    if not m:
        return None
    letter = m.group(1).upper()
    body = m.group(2).strip()
    
    vi_match = re.search(r"\nVI:\s*(.+)$", body, re.DOTALL | re.IGNORECASE)
    if vi_match:
        en_text = body[:vi_match.start()].strip()
        vi_text = vi_match.group(1).strip()
    else:
        slash_match = re.search(r"\s*/\s*(.+)$", body)
        if slash_match:
            en_text = body[:slash_match.start()].strip()
            vi_text = slash_match.group(1).strip()
        else:
            en_text = body
            vi_text = ""
    return {"id": letter, "en": en_text, "vi": vi_text}

def parse_question_header(txt: str):
    txt = txt.strip()
    parts = re.split(r"\s*[·]\s*", txt)
    if len(parts) < 2:
        return None
    qid_raw = parts[0].strip()
    if not re.match(r"^Q\d+$", qid_raw, re.IGNORECASE):
        return None
    qid = f"Q{int(qid_raw[1:]):03d}"
    standard = parts[1].strip() if len(parts) > 1 else ""
    topic_raw = parts[2].strip() if len(parts) > 2 else ""
    topic_parts = re.split(r"\s*/\s*", topic_raw, maxsplit=1)
    topic_en = topic_parts[0].strip()
    topic_vi = topic_parts[1].strip() if len(topic_parts) > 1 else ""
    return qid, standard, topic_en, topic_vi

def parse_answer_header(txt: str):
    txt = txt.strip()
    m = re.match(r"^(Q\d+)\s*[·\-]\s*(?:ĐÁP\s*ÁN|CORRECT(?:\s*ANSWER)?)\s+([A-D])", txt, re.IGNORECASE)
    if m:
        qid_raw = m.group(1)
        qid = f"Q{int(qid_raw[1:]):03d}"
        letter = m.group(2).upper()
        return qid, letter
    return None

# ─── ASHRAE 62.1-2022 Helpers ─────────────────────────────────────────────────
TOPIC_MAP = {
    "1": {"en": "Purpose and scope", "vi": "Mục đích và phạm vi"},
    "2": {"en": "Purpose and scope", "vi": "Mục đích và phạm vi"},
    "3": {"en": "Definitions, terminology, and airflow diagram", "vi": "Định nghĩa, thuật ngữ và sơ đồ luồng khí"},
    "4": {"en": "Outdoor air quality and survey records", "vi": "Chất lượng không khí ngoài trời và hồ sơ khảo sát"},
    "5": {"en": "Systems, equipment, moisture, filtration, classification, and recirculation", "vi": "Hệ thống, thiết bị, ẩm, lọc, phân loại và tuần hoàn"},
    "6": {"en": "VRP, IAQP, natural ventilation, exhaust, and calculation appendices", "vi": "VRP, IAQP, thông gió tự nhiên, hút thải và phụ lục tính toán"},
    "7": {"en": "Construction, startup, and IAQ verification", "vi": "Thi công, khởi động và xác minh IAQ"},
    "8": {"en": "Operation, maintenance, and change of use", "vi": "Vận hành, bảo trì và thay đổi mục đích sử dụng"},
}

def get_topic_for_source_ashrae62(src_text: str) -> dict:
    m = re.search(r"[§S]\s*(\d)", src_text)
    if m:
        sec_num = m.group(1)
        if sec_num in TOPIC_MAP:
            return TOPIC_MAP[sec_num]
    if any(x in src_text.upper() for x in ["APPENDIX", "PHỤ LỤC", "TABLE B-", "FIGURE F-", "HÌNH B-", "HÌNH F-"]):
        return TOPIC_MAP["6"]
    return {"en": "VRP, IAQP, natural ventilation, exhaust, and calculation appendices", "vi": "VRP, IAQP, thông gió tự nhiên, hút thải và phụ lục tính toán"}

def parse_ashrae62_option(raw_text: str, letter: str) -> dict:
    raw_text = raw_text.strip()
    m = re.match(r"^([A-D])[.)]\s+(.+)$", raw_text, re.DOTALL)
    if m:
        raw_text = m.group(2).strip()
    vi_match = re.search(r"\nVI:\s*(.+)$", raw_text, re.DOTALL | re.IGNORECASE)
    if vi_match:
        en_text = raw_text[:vi_match.start()].strip()
        vi_text = vi_match.group(1).strip()
    else:
        slash_match = re.search(r"\s*/\s*(.+)$", raw_text)
        if slash_match:
            en_text = raw_text[:slash_match.start()].strip()
            vi_text = slash_match.group(1).strip()
        else:
            en_text = raw_text
            vi_text = ""
    return {"id": letter, "en": en_text, "vi": vi_text}

def parse_ashrae62_stem(raw_text: str) -> dict:
    vi_match = re.search(r"\nVI:\s*(.+)$", raw_text, re.DOTALL | re.IGNORECASE)
    if vi_match:
        en_text = raw_text[:vi_match.start()].strip()
        vi_text = vi_match.group(1).strip()
    else:
        en_text = raw_text
        vi_text = ""
    return {"en": en_text, "vi": vi_text}

def parse_ashrae62_explanation(raw_text: str) -> dict:
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    en_lines = []
    vi_lines = []
    is_vi = False
    for line in lines:
        if line.startswith("[Reference:") or line.startswith("Reference:"):
            continue
        if line.startswith("Đáp án") or line.startswith("Giải thích:") or line.startswith("VI:") or line.startswith("Lời giải"):
            is_vi = True
        if is_vi:
            vi_lines.append(line)
        else:
            en_lines.append(line)
    en_text = "\n".join(en_lines).strip()
    vi_text = "\n".join(vi_lines).strip() if vi_lines else en_text
    if not en_text and vi_text:
        en_text = vi_text
    return {"en": en_text, "vi": vi_text}

def parse_ashrae62_2022_docx(doc_path: Path, start_qid_num=305, start_img_num=14):
    doc = Document(str(doc_path))
    warnings = []
    
    img_map, img_blobs = extract_all_images_ashrae62(doc, start_num=start_img_num)
    
    ans_table = {}
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        for r in range(1, len(t2.rows)):
            cells = [c.text.strip() for c in t2.rows[r].cells]
            for i in range(0, len(cells), 2):
                if i+1 < len(cells) and cells[i].isdigit():
                    ans_table[int(cells[i])] = cells[i+1]
                    
    paras = doc.paragraphs
    questions_raw = []
    current_bloom = "Remember"
    pending_imgs = []
    in_part_i = True
    
    for p in paras:
        style = p.style.name
        txt = p.text.strip()
        rids = get_para_image_rids(p)
        
        if style.startswith("Heading 1") and ("PART II" in txt.upper() or "PHẦN II" in txt.upper()):
            in_part_i = False
            
        if not in_part_i:
            continue
            
        for r in rids:
            if r in img_map and img_map[r] not in pending_imgs:
                pending_imgs.append(img_map[r])
                
        if style.startswith("Heading 2"):
            if "REMEMBER" in txt or "NHỚ" in txt: current_bloom = "Remember"
            elif "UNDERSTAND" in txt or "HIỂU" in txt: current_bloom = "Understand"
            elif "APPLY" in txt or "VẬN DỤNG" in txt: current_bloom = "Apply"
            elif "ANALYZE" in txt or "PHÂN TÍCH" in txt: current_bloom = "Analyze"
            elif "EVALUATE" in txt or "ĐÁNH GIÁ" in txt: current_bloom = "Evaluate"
            elif "CREATE" in txt or "SÁNG TẠO" in txt: current_bloom = "Create"
        elif style == "Question" and txt:
            q_num = len(questions_raw) + 1
            qid = f"Q{start_qid_num + q_num - 1:03d}"
            
            imgs = list(pending_imgs)
            # Figure for Questions 45-46 is shared between Q45 and Q46
            if q_num == 46 and not imgs and questions_raw and len(questions_raw[-1]["images"]) > 0:
                imgs.append(questions_raw[-1]["images"][0])
            pending_imgs.clear()
            
            q_obj = {
                "id": qid,
                "order": start_qid_num + q_num - 1,
                "sectionId": "temp",
                "standard": "ASHRAE 62.1-2022",
                "bloomLevel": current_bloom,
                "topic": {"en": "", "vi": ""},
                "stem": parse_ashrae62_stem(txt),
                "options": [],
                "correctOptionId": ans_table.get(q_num, ""),
                "explanation": {"en": "", "vi": ""},
                "sourceText": "",
                "images": imgs,
            }
            questions_raw.append(q_obj)
        elif style == "Option" and txt and questions_raw:
            cur_q = questions_raw[-1]
            letter = "ABCD"[len(cur_q["options"])] if len(cur_q["options"]) < 4 else "X"
            cur_q["options"].append(parse_ashrae62_option(txt, letter))
        elif style == "Source Ref" and txt and questions_raw:
            cur_q = questions_raw[-1]
            cur_q["sourceText"] = txt
            cur_q["topic"] = get_topic_for_source_ashrae62(txt)
        elif style == "Normal" and txt and not rids and questions_raw and len(questions_raw[-1]["options"]) == 0:
            cur_q = questions_raw[-1]
            extra = parse_ashrae62_stem(txt)
            if extra["en"]: cur_q["stem"]["en"] += "\n" + extra["en"]
            if extra["vi"]: cur_q["stem"]["vi"] += "\n" + extra["vi"]
            
    explanations = []
    for p in paras:
        if p.style.name == "Explanation" and p.text.strip():
            explanations.append(parse_ashrae62_explanation(p.text.strip()))
            
    for idx, q in enumerate(questions_raw):
        if idx < len(explanations):
            q["explanation"] = explanations[idx]
        if not q["topic"]["en"]:
            q["topic"] = get_topic_for_source_ashrae62(q["sourceText"])
            
    return questions_raw, warnings, img_blobs

def parse_ashrae522_2017_docx(doc_path: Path, start_qid_num=380, start_img_num=19):
    doc = Document(str(doc_path))
    warnings = []
    
    img_map, img_blobs = extract_all_images_ashrae62(doc, start_num=start_img_num)
    print(f"  Đã trích xuất {len(img_map)} ảnh cho ASHRAE 52.2-2017")
    
    part_ii_idx = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if "PART II" in p.text.strip() or "PHẦN II" in p.text.strip():
            part_ii_idx = i
            break
            
    answers_map = {}
    current_q_num = None
    cur_exp_lines = []
    cur_source = ""
    in_explanation = False
    
    def save_cur_ans():
        if current_q_num and current_q_num in answers_map:
            exp_text = "\n".join(cur_exp_lines).strip()
            en_exp = exp_text
            vi_exp = exp_text
            if "Giải thích:" in exp_text:
                parts = exp_text.split("Giải thích:", 1)
                en_exp = re.sub(r"^Explanation:\s*", "", parts[0].strip())
                vi_exp = parts[1].strip()
            elif "\n" in exp_text:
                parts = exp_text.split("\n", 1)
                en_exp = re.sub(r"^Explanation:\s*", "", parts[0].strip())
                vi_exp = parts[1].strip()
            else:
                en_exp = re.sub(r"^Explanation:\s*", "", en_exp)
                
            answers_map[current_q_num]["explanation"] = {"en": en_exp, "vi": vi_exp}
            if cur_source:
                answers_map[current_q_num]["sourceText"] = re.sub(r"^Source / Nguồn:\s*", "", cur_source).strip()
            elif "sourceText" not in answers_map[current_q_num]:
                answers_map[current_q_num]["sourceText"] = "ANSI/ASHRAE Standard 52.2-2017"

    for p in doc.paragraphs[part_ii_idx:]:
        st = p.style.name
        txt = p.text.strip()
        if not txt:
            continue
        if st == "Quiz Answer" or re.search(r"Câu\s+(\d+)\.\s*Đáp án:", txt) or re.search(r"Answer:\s*([ABCD])", txt):
            save_cur_ans()
            cur_exp_lines = []
            cur_source = ""
            in_explanation = True
            
            lines = [l.strip() for l in txt.split("\n") if l.strip()]
            q_num = None
            ans_letter = None
            topic_en = "ASHRAE 52.2-2017"
            topic_vi = "ASHRAE 52.2-2017"
            
            for line in lines:
                m1 = re.search(r"Câu\s+(\d+)\.\s*Đáp án:\s*([ABCD])(?:\s*\|\s*(.*))?", line, re.IGNORECASE)
                if m1:
                    q_num = int(m1.group(1))
                    ans_letter = m1.group(2).upper()
                    if m1.group(3):
                        topic_vi = m1.group(3).strip()
                else:
                    m2 = re.search(r"Answer:\s*([ABCD])(?:\s*\|\s*(.*))?", line, re.IGNORECASE)
                    if m2:
                        ans_letter = m2.group(1).upper()
                        if m2.group(2):
                            topic_en = m2.group(2).strip()
                            
            if q_num and ans_letter:
                current_q_num = q_num
                answers_map[q_num] = {
                    "correctOptionId": ans_letter,
                    "topic": {"en": topic_en, "vi": topic_vi}
                }
        elif st == "Quiz Source" or txt.startswith("Source / Nguồn:"):
            cur_source = txt
            in_explanation = False
        elif current_q_num and in_explanation and (st == "Normal" or txt.startswith("Explanation:") or txt.startswith("Giải thích:")):
            cur_exp_lines.append(txt)
            
    save_cur_ans()
    
    questions_list = []
    cur_bloom = "Remember"
    cur_q = None
    pending_images = []
    
    for i, p in enumerate(doc.paragraphs[:part_ii_idx]):
        st = p.style.name
        txt = p.text.strip()
        rids = get_para_image_rids(p)
        for rid in rids:
            if rid in img_map and img_map[rid] not in pending_images:
                pending_images.append(img_map[rid])
                
        if not txt and not rids:
            continue
            
        if "Heading" in st:
            if "Level 1" in txt or "Remember" in txt or "Nhớ" in txt:
                cur_bloom = "Remember"
            elif "Level 2" in txt or "Understand" in txt or "Hiểu" in txt:
                cur_bloom = "Understand"
            elif "Level 3" in txt or "Apply" in txt or "Vận dụng" in txt:
                cur_bloom = "Apply"
            elif "Level 4" in txt or "Analyze" in txt or "Phân tích" in txt:
                cur_bloom = "Analyze"
            elif "Level 5" in txt or "Evaluate" in txt or "Đánh giá" in txt:
                cur_bloom = "Evaluate"
            elif "Level 6" in txt or "Create" in txt or "Sáng tạo" in txt:
                cur_bloom = "Create"
        elif st == "Quiz Question":
            if cur_q and len(cur_q["options"]) == 4:
                questions_list.append(cur_q)
                
            lines = [l.strip() for l in txt.split("\n") if l.strip()]
            en_stem = lines[0] if len(lines) > 0 else txt
            vi_stem = lines[1] if len(lines) > 1 else en_stem
            
            q_num = len(questions_list) + 1
            m_num = re.search(r"Câu\s+(\d+)\.", txt)
            if m_num:
                q_num = int(m_num.group(1))
                
            vi_stem = re.sub(r"^Câu\s+\d+\.\s*", "", vi_stem).strip()
            en_stem_clean = re.sub(r"^\[.*?\]\s*", "", en_stem).strip()
            vi_stem_clean = re.sub(r"^\[.*?\]\s*", "", vi_stem).strip()
            
            global_qid_num = start_qid_num + len(questions_list)
            qid_str = f"Q{global_qid_num:03d}"
            
            cur_q = {
                "id": qid_str,
                "order": global_qid_num,
                "sectionId": "sec-temp",
                "bloomLevel": cur_bloom,
                "standard": "ASHRAE 52.2-2017",
                "stem": {"en": en_stem_clean, "vi": vi_stem_clean},
                "options": [],
                "images": [],
                "correctOptionId": "",
                "topic": {"en": "", "vi": ""},
                "sourceText": "",
                "explanation": {"en": "", "vi": ""},
                "_orig_qnum": q_num
            }
            if pending_images:
                cur_q["images"].extend(pending_images)
                pending_images = []
        elif st == "Quiz Option" and cur_q is not None:
            lines = [l.strip() for l in txt.split("\n") if l.strip()]
            if len(lines) >= 2:
                en_opt = lines[0]
                vi_opt = lines[1]
                opt_id_char = chr(65 + len(cur_q["options"]))
                vi_opt_clean = re.sub(r"^[ABCD]\.\s*", "", vi_opt).strip()
                en_opt_clean = re.sub(r"^[ABCD]\.\s*", "", en_opt).strip()
                cur_q["options"].append({
                    "id": opt_id_char,
                    "en": en_opt_clean,
                    "vi": vi_opt_clean
                })

    if cur_q and len(cur_q["options"]) == 4:
        questions_list.append(cur_q)

    for q in questions_list:
        orig_n = q.pop("_orig_qnum", None)
        if orig_n and orig_n in answers_map:
            ans_info = answers_map[orig_n]
            q["correctOptionId"] = ans_info["correctOptionId"]
            q["topic"] = ans_info["topic"]
            q["explanation"] = ans_info.get("explanation", {"en": "", "vi": ""})
            q["sourceText"] = ans_info.get("sourceText", "ANSI/ASHRAE Standard 52.2-2017")
        else:
            warnings.append(f"Không tìm thấy đáp án cho câu hỏi thứ {orig_n} ({q['id']})")

    return questions_list, warnings, img_blobs

# ─── Main Parse ───────────────────────────────────────────────────────────────
def parse_docx(doc_path: Path):
    doc = Document(str(doc_path))
    warnings = []
    
    print("  Trích xuất ảnh nhúng...")
    img_map, img_blobs = extract_all_images(doc)
    print(f"  Đã trích xuất {len(img_map)} ảnh")
    
    paras = doc.paragraphs
    print(f"  Tổng số paragraph: {len(paras)}")
    
    questions_raw = {}
    sections = []
    q_order_counter = 0
    
    current_bloom = ""
    current_standard = ""
    current_section_id = None
    section_counter = 0
    root_sec_id = None
    curr_bloom_sec_id = None
    
    in_questions_part = False
    
    i = 0
    while i < len(paras):
        para = paras[i]
        style = para.style.name
        txt = para_full_text(para).strip()
        img_rids = get_para_image_rids(para)
        
        if style.startswith("Heading"):
            if "PHẦN II" in txt or "ĐÁP ÁN" in txt.upper():
                in_questions_part = False
            elif "PHẦN I" in txt:
                in_questions_part = True
            
            bm = BLOOM_H_RE.search(txt)
            if bm:
                bloom_en  = bm.group(2).capitalize()
                current_bloom = f"{bloom_en}"
            
            sm = STD_H_RE.search(txt)
            if sm and style == "Heading 2":
                current_standard = sm.group(1).strip()
            
            if in_questions_part and txt:
                section_counter += 1
                sec_id = f"sec-{section_counter:03d}"
                
                if "PHẦN I" in txt or section_counter == 1:
                    lvl = 0
                    parent_id = None
                    root_sec_id = sec_id
                elif bm or "BLOOM" in txt.upper() or style == "Heading 1":
                    lvl = 1
                    parent_id = root_sec_id
                    curr_bloom_sec_id = sec_id
                else:
                    lvl = 2
                    parent_id = curr_bloom_sec_id
                
                sec = {
                    "id": sec_id,
                    "order": section_counter,
                    "titleEn": txt,
                    "titleVi": txt,
                    "level": lvl,
                    "bloomLevel": current_bloom if lvl > 0 else None,
                    "standard": current_standard if lvl == 2 else None,
                    "parentId": parent_id,
                    "questionIds": [],
                }
                sections.append(sec)
                current_section_id = sec["id"]
            
            i += 1
            continue
        
        if style == "Question Header" and txt:
            parsed = parse_question_header(txt)
            if parsed:
                qid, std, topic_en, topic_vi = parsed
                q_order_counter += 1
                questions_raw[qid] = {
                    "id": qid,
                    "order": q_order_counter,
                    "sectionId": current_section_id or "sec-000",
                    "standard": std or current_standard,
                    "bloomLevel": current_bloom,
                    "topic": {"en": topic_en, "vi": topic_vi},
                    "stem": {"en": "", "vi": ""},
                    "options": [],
                    "correctOptionId": "",
                    "explanation": {"en": "", "vi": ""},
                    "sourceText": "",
                    "images": [],
                    "_cur_qid": qid,
                }
                if current_section_id:
                    for sec in sections:
                        if sec["id"] == current_section_id:
                            sec["questionIds"].append(qid)
                            break
                for rid in img_rids:
                    if rid in img_map:
                        questions_raw[qid]["images"].append(img_map[rid])
            i += 1
            continue
        
        if style == "Stem English" and txt:
            last_qid = f"Q{q_order_counter:03d}"
            if last_qid in questions_raw:
                questions_raw[last_qid]["stem"]["en"] = txt
                for rid in img_rids:
                    if rid in img_map:
                        wp = img_map[rid]
                        if wp not in questions_raw[last_qid]["images"]:
                            questions_raw[last_qid]["images"].append(wp)
            i += 1
            continue
        
        if style == "Stem Vietnamese" and txt:
            last_qid = f"Q{q_order_counter:03d}"
            if last_qid in questions_raw:
                questions_raw[last_qid]["stem"]["vi"] = txt
                for rid in img_rids:
                    if rid in img_map:
                        wp = img_map[rid]
                        if wp not in questions_raw[last_qid]["images"]:
                            questions_raw[last_qid]["images"].append(wp)
            i += 1
            continue
        
        if style == "Figure Caption" and txt:
            last_qid = f"Q{q_order_counter:03d}"
            if last_qid in questions_raw:
                for rid in img_rids:
                    if rid in img_map:
                        wp = img_map[rid]
                        if wp not in questions_raw[last_qid]["images"]:
                            questions_raw[last_qid]["images"].append(wp)
            if i > 0:
                prev_para = paras[i-1]
                prev_rids = get_para_image_rids(prev_para)
                if prev_rids and last_qid in questions_raw:
                    for rid in prev_rids:
                        if rid in img_map:
                            wp = img_map[rid]
                            if wp not in questions_raw[last_qid]["images"]:
                                questions_raw[last_qid]["images"].append(wp)
            i += 1
            continue
        
        if style == "Option" and txt:
            last_qid = f"Q{q_order_counter:03d}"
            if last_qid in questions_raw:
                opt = parse_option(txt)
                if opt:
                    existing = [o for o in questions_raw[last_qid]["options"] if o["id"] == opt["id"]]
                    if not existing:
                        questions_raw[last_qid]["options"].append(opt)
                    else:
                        warnings.append(f"{last_qid}: Trùng option {opt['id']}")
                else:
                    warnings.append(f"{last_qid}: Không parse được option: {repr(txt[:80])}")
            i += 1
            continue
        
        if style == "Normal" and img_rids:
            last_qid = f"Q{q_order_counter:03d}"
            if last_qid in questions_raw:
                for rid in img_rids:
                    if rid in img_map:
                        wp = img_map[rid]
                        if wp not in questions_raw[last_qid]["images"]:
                            questions_raw[last_qid]["images"].append(wp)
            i += 1
            continue
        
        i += 1
    
    print("  Đọc phần II (đáp án và giải thích)...")
    in_answers_part = False
    current_ans_qid = None
    
    i = 0
    while i < len(paras):
        para = paras[i]
        style = para.style.name
        txt = para_full_text(para).strip()
        
        if style.startswith("Heading"):
            if "PHẦN II" in txt or "ĐÁP ÁN" in txt.upper():
                in_answers_part = True
        
        if not in_answers_part:
            i += 1
            continue
        
        if style == "Answer Header" and txt:
            parsed_ah = parse_answer_header(txt)
            if parsed_ah:
                qid, letter = parsed_ah
                current_ans_qid = qid
                if qid in questions_raw:
                    questions_raw[qid]["correctOptionId"] = letter
                else:
                    warnings.append(f"Answer Header: {qid} không có trong danh sách câu hỏi")
            i += 1
            continue
        
        if style == "Answer Text" and txt and current_ans_qid:
            i += 1
            continue
        
        if style == "Explanation" and txt and current_ans_qid:
            if current_ans_qid in questions_raw:
                m_expl = EXPL_RE.match(txt)
                if m_expl:
                    expl_text = m_expl.group(1).strip()
                else:
                    expl_text = txt
                questions_raw[current_ans_qid]["explanation"]["vi"] = expl_text
                if not questions_raw[current_ans_qid]["explanation"]["en"]:
                    questions_raw[current_ans_qid]["explanation"]["en"] = expl_text
            i += 1
            continue
        
        if style == "Source Line" and txt and current_ans_qid:
            if current_ans_qid in questions_raw:
                m_src = SRC_RE.match(txt)
                if m_src:
                    questions_raw[current_ans_qid]["sourceText"] = m_src.group(1).strip()
                else:
                    questions_raw[current_ans_qid]["sourceText"] = txt
            i += 1
            continue
        
        i += 1
    
    questions = sorted(questions_raw.values(), key=lambda q: q["order"])
    for q in questions:
        q.pop("_cur_qid", None)
    
    return sections, questions, warnings, img_map, img_blobs

def merge_ashrae_sections(sections, q2, q3, all_questions):
    blooms = ["Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create"]
    for b in blooms:
        b_621_qs = [q for q in q2 if q["bloomLevel"] == b]
        b_522_qs = [q for q in q3 if q["bloomLevel"] == b]
        l1_secs = [s for s in sections if s["level"] == 1 and b.lower() in s["titleEn"].lower()]
        if l1_secs:
            l1_sec = l1_secs[0]
            
            # Insert ASHRAE 52.2-2017 right after ASHRAE 52.2
            new_sec_522 = {
                "id": f"temp_522_{b}",
                "order": 0,
                "titleEn": f"ASHRAE 52.2-2017 · {len(b_522_qs)} câu",
                "titleVi": f"ASHRAE 52.2-2017 · {len(b_522_qs)} câu",
                "level": 2,
                "bloomLevel": None,
                "standard": "ASHRAE 52.2-2017",
                "parentId": l1_sec["id"],
                "questionIds": [q["id"] for q in b_522_qs],
            }
            insert_idx = len(sections) - 1
            for idx, s in enumerate(sections):
                if s["level"] == 2 and s["parentId"] == l1_sec["id"] and "52.2" in str(s.get("standard", "")):
                    insert_idx = idx + 1
                    break
            sections.insert(insert_idx, new_sec_522)
            
            # Insert ASHRAE 62.1-2022 right after ASHRAE 62.1-2013
            new_sec_621 = {
                "id": f"temp_621_{b}",
                "order": 0,
                "titleEn": f"ASHRAE 62.1-2022 · {len(b_621_qs)} câu",
                "titleVi": f"ASHRAE 62.1-2022 · {len(b_621_qs)} câu",
                "level": 2,
                "bloomLevel": None,
                "standard": "ASHRAE 62.1-2022",
                "parentId": l1_sec["id"],
                "questionIds": [q["id"] for q in b_621_qs],
            }
            insert_idx = len(sections) - 1
            for idx, s in enumerate(sections):
                if s["level"] == 2 and s["parentId"] == l1_sec["id"] and "62.1-2013" in str(s.get("standard", "")):
                    insert_idx = idx + 1
                    break
            sections.insert(insert_idx, new_sec_621)

    # Re-index all section IDs and Orders
    id_map = {}
    for idx, s in enumerate(sections, 1):
        id_map[s["id"]] = f"sec-{idx:03d}"
        
    for idx, s in enumerate(sections, 1):
        s["id"] = id_map[s["id"]]
        s["order"] = idx
        if s["parentId"] and s["parentId"] in id_map:
            s["parentId"] = id_map[s["parentId"]]

    for s in sections:
        if s["level"] == 2 and s["standard"]:
            for qid in s["questionIds"]:
                for q in all_questions:
                    if q["id"] == qid:
                        q["sectionId"] = s["id"]
                        break

# ─── Validation ───────────────────────────────────────────────────────────────
def validate(questions, sections, warnings, img_blobs):
    errors = []
    
    if len(questions) != EXPECTED_COUNT:
        errors.append(
            f"Số câu không đúng: mong đợi {EXPECTED_COUNT}, thực tế {len(questions)}"
        )
    
    ids = [q["id"] for q in questions]
    dup = [qid for qid in set(ids) if ids.count(qid) > 1]
    if dup:
        errors.append(f"ID trùng nhau: {dup}")
        
    root_secs = [s for s in sections if s["level"] == 0]
    bloom_secs = [s for s in sections if s["level"] == 1]
    std_secs = [s for s in sections if s["level"] == 2]
    
    if len(root_secs) != 1:
        errors.append(f"Cấu trúc section: mong đợi 1 Root level 0, thực tế {len(root_secs)}")
    if len(bloom_secs) != 6:
        errors.append(f"Cấu trúc section: mong đợi 6 Bloom level 1, thực tế {len(bloom_secs)}")
    if len(std_secs) != 36:
        errors.append(f"Cấu trúc section: mong đợi 36 Standard level 2, thực tế {len(std_secs)}")
    if len(sections) != 43:
        errors.append(f"Cấu trúc section: mong đợi tổng 43 section, thực tế {len(sections)}")
        
    for s in sections:
        if s["level"] == 0 and s["parentId"] is not None:
            errors.append(f"Section {s['id']} (level 0) phải có parentId=null")
        if s["level"] == 1 and (not s["parentId"] or s["parentId"] != root_secs[0]["id"]):
            errors.append(f"Section {s['id']} (level 1) có parentId sai: {s['parentId']}")
        if s["level"] == 2 and not any(b["id"] == s["parentId"] for b in bloom_secs):
            errors.append(f"Section {s['id']} (level 2) có parentId không trỏ đến Bloom: {s['parentId']}")

    bloom_counts = {}
    std_counts = {}
    for q in questions:
        b = q["bloomLevel"]
        std = q["standard"]
        bloom_counts[b] = bloom_counts.get(b, 0) + 1
        std_counts[std] = std_counts.get(std, 0) + 1
        
    expected_blooms = {
        "Remember": 107, "Understand": 91, "Apply": 97,
        "Analyze": 71, "Evaluate": 47, "Create": 26
    }
    for b, count in expected_blooms.items():
        if bloom_counts.get(b, 0) != count:
            errors.append(f"Số câu Bloom '{b}' sai: mong đợi {count}, thực tế {bloom_counts.get(b, 0)}")

    expected_stds = {
        "ASHRAE 55-2023": 65, "ASHRAE 52.2": 70, "ASHRAE 52.2-2017": 60,
        "ASHRAE 62.1-2013": 35, "ASHRAE 62.1-2022": 75,
        "ASHRAE 90.1-2022": 134
    }
    for std, count in expected_stds.items():
        if std_counts.get(std, 0) != count:
            errors.append(f"Số câu Tiêu chuẩn '{std}' sai: mong đợi {count}, thực tế {std_counts.get(std, 0)}")

    if len(img_blobs) != 21:
        errors.append(f"Số file ảnh trích xuất không đúng: mong đợi 21, thực tế {len(img_blobs)}")
    
    total_img_refs = sum(len(q["images"]) for q in questions)
    if total_img_refs != 53:
        errors.append(f"Tổng số liên kết ảnh–câu không đúng: mong đợi 53, thực tế {total_img_refs}")

    used_images = set()
    for q in questions:
        qid = q["id"]
        for img_path in q["images"]:
            used_images.add(img_path)
            if "img_0001.png" in img_path or "img_0001.jpg" in img_path:
                errors.append(f"{qid}: Ảnh bìa img_0001 bị gắn vào câu hỏi!")
            if img_path not in img_blobs:
                errors.append(f"{qid}: Ảnh bị gắn không tồn tại trong tập trích xuất: {img_path}")

    if len(used_images) != 20:
        errors.append(f"Số ảnh nội dung được dùng trong câu hỏi không đúng: mong đợi 20, thực tế {len(used_images)}")

    for q in questions:
        qid = q["id"]
        opt_ids = [o["id"] for o in q["options"]]
        
        if len(q["options"]) != 4:
            errors.append(f"{qid}: Số phương án = {len(q['options'])} (phải 4): {opt_ids}")
        
        for o in q["options"]:
            if not o["id"] in ["A", "B", "C", "D"]:
                errors.append(f"{qid}: Phương án không hợp lệ {o['id']}")
            if not o.get("en") or not o.get("vi"):
                errors.append(f"{qid}: Phương án {o['id']} thiếu tiếng Anh hoặc tiếng Việt")
        
        for letter in ["A", "B", "C", "D"]:
            if letter not in opt_ids:
                errors.append(f"{qid}: Thiếu phương án {letter}")
        
        if not q["correctOptionId"]:
            errors.append(f"{qid}: Không có đáp án đúng")
        elif q["correctOptionId"] not in opt_ids:
            errors.append(f"{qid}: correctOptionId='{q['correctOptionId']}' không tồn tại: {opt_ids}")
        
        if not q["explanation"]["en"] or not q["explanation"]["vi"]:
            errors.append(f"{qid}: Thiếu lời giải thích (phải đủ EN và VI)")
        
        if not q["stem"]["en"] or not q["stem"]["vi"]:
            errors.append(f"{qid}: Thiếu nội dung stem câu hỏi (phải đủ EN và VI)")
        elif not q["stem"]["en"]:
            warnings.append(f"{qid}: Thiếu stem EN")
        elif not q["stem"]["vi"]:
            warnings.append(f"{qid}: Thiếu stem VI")
    
    return errors

# ─── Build Report ─────────────────────────────────────────────────────────────
def build_report(docx_path, docx2_path, docx3_path, sections, questions, warnings, errors, img_count):
    sha1 = sha256_file(docx_path)
    sha2 = sha256_file(docx2_path)
    sha3 = sha256_file(docx3_path)
    now = datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z")
    
    sec_dist = {}
    for s in sections:
        sec_dist[s["id"]] = {"title": s["titleEn"], "count": len(s["questionIds"])}
    
    bloom_dist = {}
    for q in questions:
        b = q["bloomLevel"] or "Unknown"
        bloom_dist[b] = bloom_dist.get(b, 0) + 1
    
    std_dist = {}
    for q in questions:
        s = q["standard"] or "Unknown"
        std_dist[s] = std_dist.get(s, 0) + 1
    
    return {
        "sourceFile": f"{docx_path.name}, {docx2_path.name}, {docx3_path.name}",
        "sha256": f"{sha1}_{sha2}_{sha3}",
        "importedAt": now,
        "questionCount": len(questions),
        "sectionCount": len(sections),
        "imageCount": img_count,
        "distributionBySection": sec_dist,
        "distributionByBloom": bloom_dist,
        "distributionByStandard": std_dist,
        "warnings": warnings,
        "errors": errors,
        "status": "OK" if not errors else "FAILED",
    }

# ─── Safe JSON write ──────────────────────────────────────────────────────────
def safe_write_json(path: Path, data):
    tmp = path.with_suffix(path.suffix + ".tmp")
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    if path.exists():
        backup = path.with_suffix(path.suffix + ".bak")
        path.replace(backup)
    tmp.replace(path)

# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    print("=" * 70)
    print("HVAC Quiz – Import DOCX → questions.json")
    print("=" * 70)
    print(f"File 1     : {DOCX_PATH}")
    print(f"File 2     : {DOCX2_PATH}")
    print(f"File 3     : {DOCX3_PATH}")
    print(f"Thư mục ảnh: {IMG_DIR}")
    print(f"Output     : {QUESTIONS_JSON}")
    print()
    
    print("Bước 1: Phân tích file Word...")
    try:
        sections, questions, warnings, img_map, img_blobs = parse_docx(DOCX_PATH)
        print("  Phân tích file Word thứ 2 (ASHRAE 62.1-2022)...")
        q2, warnings2, img_blobs2 = parse_ashrae62_2022_docx(DOCX2_PATH, start_qid_num=len(questions) + 1, start_img_num=len(img_blobs) + 1)
        questions.extend(q2)
        warnings.extend(warnings2)
        img_blobs.update(img_blobs2)
        print("  Phân tích file Word thứ 3 (ASHRAE 52.2-2017)...")
        q3, warnings3, img_blobs3 = parse_ashrae522_2017_docx(DOCX3_PATH, start_qid_num=len(questions) + 1, start_img_num=len(img_blobs) + 1)
        questions.extend(q3)
        warnings.extend(warnings3)
        img_blobs.update(img_blobs3)
        merge_ashrae_sections(sections, q2, q3, questions)
    except Exception as e:
        print(f"ERROR: {e}")
        traceback.print_exc()
        sys.exit(1)
    
    print(f"  ✓ {len(sections)} phần")
    print(f"  ✓ {len(questions)} câu hỏi")
    print(f"  ✓ {len(img_blobs)} ảnh")
    print()
    
    print("Bước 2: Kiểm tra dữ liệu...")
    errors = validate(questions, sections, warnings, img_blobs)
    
    ok_count = sum(1 for q in questions if len(q["options"]) == 4 and q["correctOptionId"])
    print(f"  Câu hỏi đạt yêu cầu : {ok_count}/{len(questions)}")
    
    if errors:
        print(f"  LỖI ({len(errors)}):")
        for e in errors[:30]:
            print(f"    ✗ {e}")
        if len(errors) > 30:
            print(f"    ... và {len(errors) - 30} lỗi khác")
    else:
        print("  ✓ Không có lỗi nghiêm trọng")
    
    if warnings:
        print(f"  CẢNH BÁO ({min(len(warnings), 20)}/{len(warnings)}):")
        for w in warnings[:20]:
            print(f"    ⚠ {w}")
    print()
    
    print("Bước 3: Ghi báo cáo nhập dữ liệu...")
    report = build_report(DOCX_PATH, DOCX2_PATH, DOCX3_PATH, sections, questions, warnings, errors, len(img_blobs))
    safe_write_json(REPORT_JSON, report)
    print(f"  ✓ {REPORT_JSON}")
    
    print("=" * 70)
    if errors:
        print(f"KẾT QUẢ: THẤT BẠI ({len(errors)} lỗi)")
        print("  Không ghi đè questions.json hay xóa/thay đổi ảnh do kiểm tra thất bại.")
        print("  Xem chi tiết: data/import-report.json")
        sys.exit(1)
        
    print("Bước 4: Lưu hình ảnh và dữ liệu hợp lệ...")
    valid_fnames = set()
    for web_path, (fname, blob) in img_blobs.items():
        dest = IMG_DIR / fname
        with open(dest, "wb") as f:
            f.write(blob)
        valid_fnames.add(fname)
        
    for fpath in IMG_DIR.glob("*.*"):
        if fpath.is_file() and fpath.name not in valid_fnames:
            try:
                fpath.unlink()
                print(f"  Đã xóa ảnh cũ không sử dụng: {fpath.name}")
            except Exception as e:
                print(f"  Cảnh báo: không xóa được {fpath.name}: {e}")
                
    sha1 = sha256_file(DOCX_PATH)
    sha2 = sha256_file(DOCX2_PATH)
    sha3 = sha256_file(DOCX3_PATH)
    now = datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z")
    
    output = {
        "schemaVersion": 1,
        "source": {
            "fileName": f"{DOCX_PATH.name}, {DOCX2_PATH.name}, {DOCX3_PATH.name}",
            "sha256": f"{sha1}_{sha2}_{sha3}",
            "importedAt": now,
            "questionCount": len(questions),
        },
        "sections": sections,
        "questions": questions,
    }
    
    safe_write_json(QUESTIONS_JSON, output)
    print(f"  ✓ {QUESTIONS_JSON}")
    print()
    
    print("KẾT QUẢ: THÀNH CÔNG")
    print(f"  Câu hỏi : {len(questions)}")
    print(f"  Phần    : {len(sections)}")
    print(f"  Ảnh     : {len(img_blobs)}")
    print()
    print("Chạy tiếp:")
    print("  npm run dev")
    print("=" * 70)

if __name__ == "__main__":
    main()
